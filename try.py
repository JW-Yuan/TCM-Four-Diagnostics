from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import datetime

def generate_tcm_report(data, output_docx="tcm_report.docx", output_pdf="tcm_report.pdf"):
    doc = Document()

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    def add_key_value(key, value):
        para = doc.add_paragraph()
        para.add_run(f"{key}：").bold = True
        para.add_run(str(value))

    def add_table(title, content_dict):
        add_heading(title)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目名称'
        hdr_cells[1].text = '检查结果'
        for key, value in content_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        doc.add_paragraph()

    # 基本信息
    add_heading("中医四诊报告")
    for key in ["姓名", "性别", "年龄", "既往病史", "检测时间"]:
        add_key_value(key, data.get(key, ""))

    doc.add_paragraph("机器人智能传感与自主控制实验室")
    doc.add_paragraph()

    # 各诊断信息
    for section in ["舌诊", "面诊", "脉诊"]:
        add_key_value("姓名", data["姓名"])
        add_key_value("性别", data["性别"])
        add_key_value("年龄", data["年龄"])
        add_table(section, data[section]["项目"])
        add_key_value("检查时间", data[section]["检查时间"])
        doc.add_paragraph()

    # 问诊
    add_key_value("姓名", data["姓名"])
    add_key_value("性别", data["性别"])
    add_key_value("年龄", data["年龄"])
    add_heading("问诊")
    for key in ["症状", "病机分析", "调理建议", "就医提示"]:
        add_key_value(key, data["问诊"][key])
    add_key_value("检查时间", data["问诊"]["检查时间"])

    # 保存 Word 再转 PDF
    doc.save(output_docx)
    convert(output_docx, output_pdf)

# 示例数据传入
data = {
    "姓名": "匿名",
    "性别": "女",
    "年龄": 28,
    "既往病史": "无",
    "检测时间": "2025-04-20",
    "舌诊": {
        "项目": {
            "舌像": "",
            "舌神": "荣舌",
            "舌色": "红舌",
            "舌形": "裂纹舌",
            "舌态": "正常舌",
            "苔质": "润苔",
            "苔色": "黄苔"
        },
        "检查时间": "2025-04-20 17:12:11"
    },
    "面诊": {
        "项目": {
            "面像": "",
            "左颊": "偏黑",
            "右颊": "偏黄",
            "颌": "偏黑",
            "鼻": "偏黄",
            "庭": "偏黑"
        },
        "检查时间": "2025-04-20 17:12:56"
    },
    "脉诊": {
        "项目": {
            "平均血氧饱和度": "88.45%",
            "平均脉率": "85.70bpm",
            "平均灌注指数": "15.20%",
            "脉像": "平脉"
        },
        "检查时间": "2025-04-20 17:14:06"
    },
    "问诊": {
        "症状": "失眠多梦，容易疲惫，饮食不规律",
        "病机分析": "心脾两虚，阴血不足……",
        "调理建议": "饮食建议、运动建议、中药建议……",
        "就医提示": "若持续心悸、潮热等……",
        "检查时间": "2025-04-20 17:15:27"
    }
}

generate_tcm_report(data)
